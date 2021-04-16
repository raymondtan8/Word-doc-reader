#! python3
# For my reference: python3 -i fileread.py

"""
File Reader for Edison Documents. Not really reusable for any other purposes.

Author: Raymond Tan
Created: 11/15/20

"""

"""
ERRORS THAT NEED TO BE FIXED:

Fix test function -- seems to double the actual amount for everything.

"""

import docx

data = []
raised_error_flag = False


def test(filename='demo.docx'):
    """Runs a test of the main components."""
    init_data(filename)
    for x in data:
        print(x)
    check_totals(filename)


def return_doc_object(filename):
    """Returns a doc object of FILENAME."""
    return docx.Document(filename)


def get_text(filename):
    """Returns the text in each paragraph of an entire document,
    with the text of each paragraph on its own line.
    """
    doc = return_doc_object(filename)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return '\n'.join(full_text)


def raise_error(dic, type, percent=.05, col='', actual_due=''):
    """Raises an error of type TYPE. Sets RAISED_ERROR_FLAG
    to True.
    """
    global raised_error_flag  # Need this statement to change value of a global variable
    raised_error_flag = True
    if type == 'Fee error':
        fee = dic['Fee']
        unit = dic['Unit']
        actual_fee = round(dic['Rent'] * percent, 2)
        print(
            f'Current fee is ${fee} for Unit {unit}, but should be ${actual_fee}.')
    elif type == 'Total error':
        old_total = dic[col]
        print(
            f'Current TOTAL for {col} is ${old_total}, but should be ${total_column(col)}.')
    elif type == 'Due owner error':
        old_due = dic['Due owner']
        print(f'Current due owner is ${old_due}, but should be ${actual_due}.')


def get_totals(filename):
    """Returns a dictionary with the value of the totals of each category:
    RENT, FEE, and REPAIRS.
    """
    table = return_doc_object(filename).tables[0]
    table2 = return_doc_object(filename).tables[1]
    total = {}
    # Violating the abstraction barrier
    for i, cell in enumerate(table.rows[13].cells):
        if i == 1:
            total['Rent'] = round(float(cell.text[1:].replace(',', '')), 2)
        elif i == 2:
            total['Fee'] = round(float(cell.text[1:].replace(',', '')), 2)
        elif i == 3:
            total['Repairs'] = round(float(cell.text[1:].replace(',', '')), 2)
    total['Due owner'] = round(
        float(table2.rows[0].cells[1].text[-9:].replace(',', '')), 2)
    return total


def total_column(column):
    """Returns the total value of totaling a column of type COLUMN."""
    total_val = 0
    for dic in data:
        for key in dic.keys():
            if key == column:
                total_val += dic[key]
    return round(total_val, 2)


def check_totals(filename):
    """Matches the total calculated with the actual total listed on the document.
    Raises an error if numbers do not match up.
    """
    total_on_doc = get_totals(filename)
    for col in total_on_doc.keys():
        if col == 'Due owner':
            actual_due_owner = total_on_doc['Rent'] - \
                total_on_doc['Fee'] - total_on_doc['Repairs']
            if actual_due_owner != total_on_doc['Due owner']:
                raise_error(total_on_doc, 'Due owner error',
                            actual_due=actual_due_owner)
        elif total_column(col) != total_on_doc[col]:
            raise_error(total_on_doc, 'Total error', col=col)


def init_data(filename, start=0, stop=13):
    """Each row in a table has a corresponding dictionary
    in the global list DATA. All of the data simply gets put
    into its corrsponding category.
    """
    doc = return_doc_object(filename)
    table = doc.tables[0]
    for i, row in enumerate(table.rows):
        if i == start or i == stop:
            continue

        this_dict = {}
        for j, cell in enumerate(row.cells):
            if j == 0:
                this_dict['Unit'] = i
            elif j == 1:
                this_dict['Rent'] = cell.text
            elif j == 2:
                this_dict['Fee'] = cell.text
            else:
                this_dict['Repairs'] = cell.text
        clean_data(this_dict)
        data.append(this_dict)


def clean_data(this_dict):
    """Takes the raw data from a dictionary and cleans up its values.
    Cleaning up means getting rid of any words, and replacing it with numbers.
    Also, replacing numbers represented in dollars to regular numbers.
    """
    # Rent will be the most involved to clean
    rent = this_dict['Rent']
    fee = this_dict['Fee']
    repairs = this_dict['Repairs']

    # Cleaning rent
    lines = []
    for line in rent.split('\n'):
        # Sorry for the huge violations of abstraction barrier
        # Note: This would have been so much cleaner with regular expressions but I was too lazy
        if len(line) > 0 and (line.split()[0] == 'Paid' or line[0] == '$'):
            to_extend = [float(ele[1:].replace(',', ''))
                         for ele in line.split() if ele[1].isdigit()]
            lines.extend(to_extend)
    this_dict['Rent'] = round(sum(lines), 2)

    # Cleaning fee
    new_fee = fee.split('\n')[0]  # This only gives us the dollar amount
    new_fee = float(new_fee.strip().strip('$'))
    this_dict['Fee'] = new_fee

    # Cleaning repairs
    new_repair = 0
    for i in range(repairs.count('$')):
        this_line = repairs.split('\n')[i]
        this_line = float(this_line.strip().strip('$').replace(',', ''))
        new_repair += this_line
    this_dict['Repairs'] = new_repair


def check_fee_percent(percent=.05):
    """Checks that each fee is PERCENT percent of the rent."""
    for dic in data:
        rent = dic['Rent']
        fee = dic['Fee']
        if rent == 0:
            if fee != 0:
                raise_error(dic, 'Fee error', percent)
            else:
                continue
        actual_percent = round(fee / rent, 6)
        if actual_percent != percent:
            raise_error(dic, 'Fee error', percent)


def action(filename):
    """Where the action happens. Only this function needs to be called
    with FILENAME for the program to work.
    """
    init_data(filename)
    check_fee_percent()
    check_totals(filename)
    if raised_error_flag == False:
        print('No errors found!')


# Run program
print('Welcome to Raymond\'s File Reader!')
print('Make sure that your file is in the same directory as this program, and that it is a .docx file (a Word Document).')
print('Type in the name of your file whenever you are ready:')
filename = input()
print('ERRORS:')
action(filename)
